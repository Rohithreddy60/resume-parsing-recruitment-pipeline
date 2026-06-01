"""
DOCX Resume Parser using python-docx.
Extracts text from paragraphs, tables, headers, and footers
while preserving structural information for NLP processing.
"""
import io
import logging
from pathlib import Path
from typing import Union, Optional
from dataclasses import dataclass, field

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)


class DOCXParserError(Exception):
    """Raised when DOCX parsing fails."""
    pass


@dataclass
class DocumentSection:
    """Structured representation of a resume section."""
    heading: Optional[str] = None
    level: int = 0
    content: list[str] = field(default_factory=list)
    is_table: bool = False
    table_data: list[list[str]] = field(default_factory=list)


class DOCXResumeParser:
    """
    Extracts text and structure from DOCX resumes using python-docx.

    Features:
    - Paragraph-level extraction with heading detection
    - Table extraction with row/cell structure
    - Header and footer text extraction
    - Text run concatenation with formatting awareness
    - Section boundary detection based on heading styles
    """

    # Heading styles that indicate section boundaries in resumes
    HEADING_STYLES = {
        "Heading 1", "Heading 2", "Heading 3",
        "Heading 4", "Title", "Subtitle",
    }

    def extract_text_from_bytes(self, content: bytes, filename: str = "resume.docx") -> str:
        """Extract full text from DOCX bytes."""
        try:
            doc = Document(io.BytesIO(content))
            return self._extract_full_text(doc)
        except Exception as e:
            logger.error(f"Failed to parse DOCX bytes [{filename}]: {e}")
            raise DOCXParserError(f"Cannot parse DOCX '{filename}': {e}") from e

    def extract_text_from_path(self, file_path: Union[str, Path]) -> str:
        """Extract text from a local DOCX file path."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")

        try:
            doc = Document(str(path))
            return self._extract_full_text(doc)
        except DOCXParserError:
            raise
        except Exception as e:
            logger.error(f"Failed to parse DOCX [{path}]: {e}")
            raise DOCXParserError(f"Cannot parse '{path}': {e}") from e

    def _extract_full_text(self, doc: Document) -> str:
        """Extract all text elements from a Document object."""
        parts = []

        # Extract header text (name often appears here)
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

        # Extract body paragraphs
        for para in doc.paragraphs:
            text = self._extract_paragraph_text(para)
            if text:
                parts.append(text)

        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        # Extract footer text
        for section in doc.sections:
            if section.footer:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)

        full_text = "\n".join(parts)
        logger.info(f"Extracted {len(full_text)} chars from DOCX ({len(doc.paragraphs)} paragraphs)")
        return full_text

    def _extract_paragraph_text(self, paragraph) -> str:
        """
        Extract text from a paragraph including all runs.
        Handles hyperlinks embedded in XML.
        """
        text_parts = []
        for child in paragraph._element:
            if child.tag == qn("w:r"):  # Regular text run
                for t in child.findall(qn("w:t")):
                    text_parts.append(t.text or "")
            elif child.tag == qn("w:hyperlink"):  # Hyperlink text
                for run in child.findall(f".//{qn('w:r')}"):
                    for t in run.findall(qn("w:t")):
                        text_parts.append(t.text or "")
        return "".join(text_parts).strip()

    def extract_sections(self, content: bytes) -> list[DocumentSection]:
        """
        Extract structured sections from a DOCX resume.
        Returns a list of DocumentSection objects where each section
        represents a logical division (e.g., Education, Experience, Skills).
        """
        doc = Document(io.BytesIO(content))
        sections = []
        current_section = DocumentSection()

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                para = docx.text.paragraph.Paragraph(element, doc)
                style_name = para.style.name if para.style else ""
                text = self._extract_paragraph_text(para)

                if style_name in self.HEADING_STYLES and text:
                    # Save current section if it has content
                    if current_section.content or current_section.heading:
                        sections.append(current_section)
                    level = int(style_name.split()[-1]) if style_name[0:7] == "Heading" else 0
                    current_section = DocumentSection(heading=text, level=level)
                elif text:
                    current_section.content.append(text)

            elif tag == "tbl":
                table = docx.table.Table(element, doc)
                table_data = [
                    [cell.text.strip() for cell in row.cells]
                    for row in table.rows
                ]
                current_section.content.append(
                    "\n".join(" | ".join(row) for row in table_data if any(row))
                )

        # Don't forget the last section
        if current_section.content or current_section.heading:
            sections.append(current_section)

        return sections

    def get_metadata(self, content: bytes) -> dict:
        """Extract DOCX core properties (author, created date, etc.)."""
        doc = Document(io.BytesIO(content))
        props = doc.core_properties
        return {
            "author": props.author,
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None,
            "title": props.title,
            "subject": props.subject,
            "keywords": props.keywords,
        }
